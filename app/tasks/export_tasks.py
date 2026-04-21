# DEPRECATED: Use @giga-pdf/pdf-engine via Next.js API routes instead
"""
Export Celery tasks.

Async tasks for document export to various formats.
"""

import asyncio
import io
import logging
import os
import zipfile
from datetime import datetime, timezone, timedelta
from typing import Optional
from uuid import uuid4

import pikepdf
import pdfplumber

from celery import shared_task

from app.tasks.celery_app import celery_app
from app.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()

# Export files directory
EXPORT_DIR = os.path.join(settings.storage_path, "exports")


def _ensure_export_dir():
    """Ensure export directory exists."""
    os.makedirs(EXPORT_DIR, exist_ok=True)


def _save_export_data(data: bytes, format: str, document_id: str) -> str:
    """
    Save export data to file system.

    Args:
        data: Binary data to save.
        format: File format/extension.
        document_id: Document identifier.

    Returns:
        File path where data was saved.
    """
    _ensure_export_dir()

    # Generate unique filename with timestamp
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    unique_id = str(uuid4())[:8]
    filename = f"{document_id}_{timestamp}_{unique_id}.{format}"
    filepath = os.path.join(EXPORT_DIR, filename)

    with open(filepath, "wb") as f:
        f.write(data)

    logger.info(f"Saved export file: {filepath} ({len(data)} bytes)")
    return filepath


def _get_document_bytes(document_id: str) -> Optional[bytes]:
    """
    Get document bytes from storage service.

    Attempts to load from Redis session first, then from database/storage.
    """
    import asyncio

    async def _load_from_storage():
        """Load document from storage service."""
        from app.core.database import get_db_session
        from app.models.database import StoredDocument, DocumentVersion
        from app.services.storage_service import storage_service
        from sqlalchemy import select

        async with get_db_session() as session:
            # First try to find stored document by session ID (document_id)
            # Check if this is a session ID or a stored document ID
            stmt = select(StoredDocument).where(StoredDocument.id == document_id)
            result = await session.execute(stmt)
            doc = result.scalar_one_or_none()

            if doc:
                # Get latest version
                version_stmt = select(DocumentVersion).where(
                    DocumentVersion.document_id == document_id
                ).order_by(DocumentVersion.version_number.desc()).limit(1)
                version_result = await session.execute(version_stmt)
                version = version_result.scalar_one_or_none()

                if version and version.file_path:
                    # Load from storage
                    return await storage_service.load_document(version.file_path)

        return None

    async def _load_from_session():
        """Try to load from Redis session."""
        from app.repositories.document_repo import document_sessions

        # Try session from Redis (get_session_async is an alias for get_session)
        session = await document_sessions.get_session(document_id)
        if session and session.pdf_doc:
            # Return the document bytes from session
            return session.original_bytes
        return None

    try:
        # Try Redis session first
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)

        try:
            pdf_bytes = loop.run_until_complete(_load_from_session())
            if pdf_bytes:
                logger.debug(f"Loaded document {document_id} from Redis session")
                return pdf_bytes

            # Try storage service
            pdf_bytes = loop.run_until_complete(_load_from_storage())
            if pdf_bytes:
                logger.debug(f"Loaded document {document_id} from storage")
                return pdf_bytes

        finally:
            loop.close()

    except Exception as e:
        logger.error(f"Failed to load document {document_id}: {e}")

    return None


@celery_app.task(bind=True, name="app.tasks.export_tasks.export_document")
def export_document(
    self,
    document_id: str,
    format: str,
    page_range: Optional[str] = None,
    dpi: int = 150,
    quality: int = 85,
    single_file: bool = False,
) -> dict:
    """
    Export document to various formats.

    Args:
        self: Celery task instance.
        document_id: Document identifier.
        format: Output format (png, jpeg, svg, docx, xlsx, html, txt).
        page_range: Pages to export (all if None).
        dpi: Resolution for image formats.
        quality: Quality for JPEG/WebP.
        single_file: Combine into single file/archive.

    Returns:
        dict: Export result with file_path (not binary data).
    """
    from app.core.preview import PreviewGenerator
    from app.repositories.document_repo import document_sessions
    from app.utils.helpers import parse_page_range

    logger.info(f"Starting export for document {document_id} to {format}")

    # Try to get session from local cache or Redis
    session = asyncio.run(document_sessions.get_session(document_id))

    if session:
        pdf_bytes = session.pdf_doc.tobytes()
        page_count = session.pdf_doc.page_count
    else:
        # Load from Redis/Storage
        pdf_bytes = _get_document_bytes(document_id)
        if not pdf_bytes:
            raise ValueError(f"Document not found: {document_id}")

        with pikepdf.Pdf.open(io.BytesIO(pdf_bytes)) as _pdf:
            page_count = len(_pdf.pages)

    logger.debug(
        "Export task: document=%s format=%s pages=%d", document_id, format, page_count
    )

    # Parse page range
    if page_range:
        pages = parse_page_range(page_range, page_count)
    else:
        pages = list(range(1, page_count + 1))

    total_pages = len(pages)
    exported_files = []

    # Image formats
    if format in ("png", "jpeg", "webp", "svg"):
        generator = PreviewGenerator(pdf_bytes)

        for idx, page_num in enumerate(pages):
            # Update progress
            progress = ((idx + 1) / total_pages) * 100
            self.update_state(
                state="PROGRESS",
                meta={"progress": progress, "message": f"Exporting page {page_num}"},
            )

            # Render page
            image_data = generator.render_page(
                page_number=page_num,
                dpi=dpi,
                format=format,
                quality=quality,
            )

            exported_files.append({
                "page": page_num,
                "format": format,
                "size": len(image_data),
                "data": image_data,
            })

        # Always create output (ZIP for multiple files, single file otherwise)
        if len(exported_files) == 1 and not single_file:
            # Single page, single file output
            file_path = _save_export_data(
                exported_files[0]["data"],
                format,
                document_id
            )
            return {
                "document_id": document_id,
                "format": format,
                "pages": 1,
                "file_path": file_path,
            }
        else:
            # Multiple pages or single_file=True: create ZIP
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                for file in exported_files:
                    filename = f"page_{file['page']:04d}.{format}"
                    zf.writestr(filename, file["data"])

            file_path = _save_export_data(
                zip_buffer.getvalue(),
                "zip",
                document_id
            )
            return {
                "document_id": document_id,
                "format": "zip",
                "original_format": format,
                "pages": len(exported_files),
                "file_path": file_path,
            }

    # Text format — pdfplumber replaces fitz page.get_text()
    elif format == "txt":
        text_parts = []

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as plumber_pdf:
            for idx, page_num in enumerate(pages):
                progress = ((idx + 1) / total_pages) * 100
                self.update_state(
                    state="PROGRESS",
                    meta={"progress": progress, "message": f"Extracting text from page {page_num}"},
                )

                text = plumber_pdf.pages[page_num - 1].extract_text() or ""
                text_parts.append(f"--- Page {page_num} ---\n{text}\n")

        text_content = "\n".join(text_parts).encode("utf-8")
        file_path = _save_export_data(text_content, "txt", document_id)

        return {
            "document_id": document_id,
            "format": "txt",
            "pages": len(pages),
            "file_path": file_path,
        }

    # HTML format — pdfplumber replaces fitz page.get_text("html")
    elif format == "html":
        html_parts = [
            '<!DOCTYPE html>',
            '<html lang="en">',
            '<head>',
            '<meta charset="utf-8">',
            '<meta name="viewport" content="width=device-width, initial-scale=1">',
            '<title>PDF Export</title>',
            '<style>',
            'body { font-family: system-ui, -apple-system, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }',
            '.page { margin-bottom: 40px; padding: 20px; border: 1px solid #e0e0e0; border-radius: 8px; }',
            '.page-header { color: #666; font-size: 14px; margin-bottom: 15px; border-bottom: 1px solid #eee; padding-bottom: 10px; }',
            '.page-text { white-space: pre-wrap; }',
            '</style>',
            '</head>',
            '<body>',
        ]

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as plumber_pdf:
            for idx, page_num in enumerate(pages):
                progress = ((idx + 1) / total_pages) * 100
                self.update_state(
                    state="PROGRESS",
                    meta={"progress": progress, "message": f"Converting page {page_num}"},
                )

                text = plumber_pdf.pages[page_num - 1].extract_text() or ""
                # Escape HTML special chars to prevent injection
                import html as _html_mod
                text_escaped = _html_mod.escape(text)
                html_parts.append(f'<div class="page" data-page="{page_num}">')
                html_parts.append(f'<div class="page-header">Page {page_num}</div>')
                html_parts.append(f'<div class="page-text">{text_escaped}</div>')
                html_parts.append('</div>')

        html_parts.append("</body>")
        html_parts.append("</html>")

        html_content = "\n".join(html_parts).encode("utf-8")
        file_path = _save_export_data(html_content, "html", document_id)

        return {
            "document_id": document_id,
            "format": "html",
            "pages": len(pages),
            "file_path": file_path,
        }

    # Word (DOCX) format — pdfplumber replaces fitz page.get_text("dict")
    elif format == "docx":
        from docx import Document as DocxDocument
        from docx.shared import Pt

        docx_doc = DocxDocument()

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as plumber_pdf:
            for idx, page_num in enumerate(pages):
                progress = ((idx + 1) / total_pages) * 100
                self.update_state(
                    state="PROGRESS",
                    meta={"progress": progress, "message": f"Converting page {page_num} to Word"},
                )

                # Page break between pages
                if idx > 0:
                    docx_doc.add_page_break()

                plumber_page = plumber_pdf.pages[page_num - 1]
                text = plumber_page.extract_text() or ""

                # Add each non-empty line as a paragraph
                for line in text.splitlines():
                    stripped = line.strip()
                    if stripped:
                        docx_doc.add_paragraph(stripped)

                # Extract tables via pdfplumber
                try:
                    for table in plumber_page.extract_tables():
                        for row_data in table:
                            cells = [str(c or "") for c in row_data]
                            docx_doc.add_paragraph("\t".join(cells))
                        docx_doc.add_paragraph("")  # blank line after table
                except Exception as exc:
                    logger.debug("Table extraction skipped for page %d: %s", page_num, exc)

        # Save to bytes
        docx_buffer = io.BytesIO()
        docx_doc.save(docx_buffer)
        docx_buffer.seek(0)

        file_path = _save_export_data(docx_buffer.getvalue(), "docx", document_id)

        return {
            "document_id": document_id,
            "format": "docx",
            "pages": len(pages),
            "file_path": file_path,
        }

    # Excel (XLSX) format — pdfplumber replaces fitz page.get_text() + page.find_tables()
    elif format == "xlsx":
        from openpyxl import Workbook
        from openpyxl.styles import Font

        wb = Workbook()
        ws = wb.active
        ws.title = "PDF Content"

        current_row = 1

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as plumber_pdf:
            for idx, page_num in enumerate(pages):
                progress = ((idx + 1) / total_pages) * 100
                self.update_state(
                    state="PROGRESS",
                    meta={"progress": progress, "message": f"Converting page {page_num} to Excel"},
                )

                # Page header row
                ws.cell(row=current_row, column=1, value=f"--- Page {page_num} ---")
                ws.cell(row=current_row, column=1).font = Font(bold=True, size=14)
                current_row += 2

                plumber_page = plumber_pdf.pages[page_num - 1]
                text = plumber_page.extract_text() or ""
                for line in text.splitlines():
                    if line.strip():
                        ws.cell(row=current_row, column=1, value=line)
                        current_row += 1

                current_row += 2  # Space between pages

                # Tables via pdfplumber
                try:
                    for table in plumber_page.extract_tables():
                        ws.cell(row=current_row, column=1, value="[Table]")
                        ws.cell(row=current_row, column=1).font = Font(italic=True)
                        current_row += 1

                        for row_data in table:
                            for col_idx, cell_value in enumerate(row_data):
                                ws.cell(
                                    row=current_row,
                                    column=col_idx + 1,
                                    value=cell_value or "",
                                )
                            current_row += 1

                        current_row += 1  # Space after table
                except Exception as exc:
                    logger.debug("Table extraction skipped for page %d: %s", page_num, exc)

        # Auto-fit column width
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except Exception:
                    pass
            ws.column_dimensions[column_letter].width = min(max_length + 2, 100)

        # Save to bytes
        xlsx_buffer = io.BytesIO()
        wb.save(xlsx_buffer)
        xlsx_buffer.seek(0)

        file_path = _save_export_data(xlsx_buffer.getvalue(), "xlsx", document_id)

        return {
            "document_id": document_id,
            "format": "xlsx",
            "pages": len(pages),
            "file_path": file_path,
        }

    else:
        raise ValueError(f"Unsupported export format: {format}")


@celery_app.task(name="app.tasks.export_tasks.cleanup_expired_exports")
def cleanup_expired_exports(max_age_hours: int = 24) -> dict:
    """
    Cleanup old export files.

    Args:
        max_age_hours: Maximum age in hours before deletion.

    Returns:
        dict: Cleanup statistics.
    """
    _ensure_export_dir()

    cutoff_time = datetime.now(timezone.utc) - timedelta(hours=max_age_hours)
    deleted_count = 0
    deleted_bytes = 0
    errors = []

    try:
        for filename in os.listdir(EXPORT_DIR):
            filepath = os.path.join(EXPORT_DIR, filename)

            try:
                # Check file modification time
                mtime = datetime.fromtimestamp(os.path.getmtime(filepath), tz=timezone.utc)

                if mtime < cutoff_time:
                    file_size = os.path.getsize(filepath)
                    os.remove(filepath)
                    deleted_count += 1
                    deleted_bytes += file_size
                    logger.debug(f"Deleted expired export: {filename}")

            except Exception as e:
                errors.append(f"{filename}: {str(e)}")
                logger.warning(f"Failed to cleanup {filename}: {e}")

    except Exception as e:
        logger.error(f"Failed to list export directory: {e}")
        errors.append(f"Directory listing: {str(e)}")

    logger.info(f"Cleanup complete: deleted {deleted_count} files ({deleted_bytes / 1024 / 1024:.2f} MB)")

    return {
        "deleted_count": deleted_count,
        "deleted_bytes": deleted_bytes,
        "errors": errors,
    }
